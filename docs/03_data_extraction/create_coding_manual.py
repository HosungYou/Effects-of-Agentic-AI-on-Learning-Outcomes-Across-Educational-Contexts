#!/usr/bin/env python3
"""
Generate Word Coding Manual for Agentic AI Learning Outcomes Meta-Analysis
Created: 2026-02-16
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def setup_styles(doc):
    """Configure document styles"""
    # Title style
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)

    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.color.rgb = RGBColor(0, 51, 102)
        if i == 1:
            heading_font.size = Pt(16)
        elif i == 2:
            heading_font.size = Pt(14)
        else:
            heading_font.size = Pt(12)

def add_title_page(doc):
    """Add title page"""
    doc.add_paragraph('Coding Manual for Agentic AI Learning Outcomes Meta-Analysis', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph('Version 1.0')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    p = doc.add_paragraph('February 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    doc.add_page_break()

def add_toc(doc):
    """Add table of contents (manual)"""
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1. Overview & Purpose', '3'),
        ('2. Research Questions', '4'),
        ('3. Inclusion/Exclusion Criteria', '5'),
        ('4. Effect Size Extraction', '7'),
        ('   4.1 Computing Hedges\' g from Means and SDs', '7'),
        ('   4.2 Converting from Other Statistics', '8'),
        ('   4.3 Handling Pre-Post Designs', '9'),
        ('   4.4 Multiple Effect Sizes per Study', '10'),
        ('   4.5 Missing Data Handling', '11'),
        ('5. AI Agent Characteristic Coding', '12'),
        ('   5.1 Human Oversight Level', '13'),
        ('   5.2 Agent Architecture', '15'),
        ('   5.3 Agency Level (APCP Framework)', '16'),
        ('   5.4 Agent Role', '18'),
        ('   5.5 Agent Modality', '19'),
        ('   5.6 AI Technology', '20'),
        ('   5.7 Adaptivity Level', '21'),
        ('6. Learning Context Coding', '22'),
        ('7. Outcome Classification', '24'),
        ('8. Quality Assessment', '26'),
        ('9. AI-Assisted Coding Protocol', '28'),
        ('10. Inter-Coder Reliability', '30'),
        ('11. Discrepancy Resolution', '31'),
        ('Appendix A: Decision Trees', '32'),
        ('Appendix B: Example Coded Study', '34'),
    ]

    table = doc.add_table(rows=len(toc_items), cols=2)
    table.style = 'Light List Accent 1'

    for idx, (item, page) in enumerate(toc_items):
        row = table.rows[idx]
        row.cells[0].text = item
        row.cells[1].text = page
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

def add_section_1(doc):
    """Section 1: Overview & Purpose"""
    doc.add_heading('1. Overview & Purpose', level=1)

    doc.add_paragraph(
        'This coding manual provides standardized procedures for extracting and coding study characteristics '
        'and effect size data for a meta-analysis examining the impact of agentic AI systems on learning outcomes '
        'across educational contexts.'
    )

    doc.add_heading('1.1 Scope', level=2)
    doc.add_paragraph(
        'This meta-analysis focuses specifically on AI systems exhibiting agency characteristics (autonomy, '
        'proactivity, adaptivity, and social ability) as defined by the APCP framework (Yan et al., 2025). '
        'Studies must report quantitative learning outcome data comparing agentic AI interventions to control conditions.'
    )

    doc.add_heading('1.2 Who Uses This Manual', level=2)
    doc.add_paragraph('Primary users:')
    items = [
        'Human coders extracting data from included studies',
        'AI coding assistants (GPT-4, Claude Opus) in the 7-phase extraction pipeline',
        'Senior reviewers reconciling discrepancies',
        'Quality assurance reviewers conducting inter-coder reliability checks'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

def add_section_2(doc):
    """Section 2: Research Questions"""
    doc.add_heading('2. Research Questions', level=1)

    rqs = [
        ('RQ1: Overall Effectiveness',
         'What is the overall effect of agentic AI systems on learning outcomes compared to control conditions?',
         'H1: Agentic AI systems will show moderate positive effects (g = 0.40-0.60) on learning outcomes.'),

        ('RQ2: Human Oversight Moderation',
         'Does the level of human oversight moderate the effectiveness of agentic AI on learning outcomes?',
         'H2: AI-led systems with checkpoints will show larger effects than fully autonomous or human-led systems.'),

        ('RQ3: Agency Level Moderation',
         'Does the level of AI agency (APCP framework) moderate learning effectiveness?',
         'H3: Higher agency levels (co-learner, peer) will show larger effects than lower levels (adaptive, proactive).'),

        ('RQ4: Learning Context Moderation',
         'Do effects vary across educational contexts (K-12, higher education, workplace) and subject domains?',
         'H4: Effects will be largest in STEM domains and higher education settings.'),

        ('RQ5: Agent Architecture',
         'Do multi-agent systems show different effectiveness compared to single-agent systems?',
         'H5: Multi-agent systems will show larger effects due to distributed expertise and peer learning dynamics.')
    ]

    for title, question, hypothesis in rqs:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.add_run('Question: ').bold = True
        p.add_run(question)
        p = doc.add_paragraph()
        p.add_run('Hypothesis: ').bold = True
        p.add_run(hypothesis)

    doc.add_page_break()

def add_section_3(doc):
    """Section 3: Inclusion/Exclusion Criteria"""
    doc.add_heading('3. Inclusion/Exclusion Criteria', level=1)

    doc.add_heading('3.1 Inclusion Criteria', level=2)
    doc.add_paragraph('Studies must meet ALL of the following criteria:')

    criteria = [
        ('Population', 'Learners in any formal or informal educational setting (K-12, higher education, workplace, professional training)'),
        ('Intervention', 'AI system exhibiting at least TWO of the four APCP agency characteristics: Autonomy, Proactivity, Cooperativity, Personalization'),
        ('Comparison', 'Non-agentic control condition (no AI, non-agentic AI, traditional instruction, or human-only instruction)'),
        ('Outcome', 'Quantitative learning outcome measure (cognitive knowledge, skills, performance, or affective outcomes)'),
        ('Study Design', 'Experimental or quasi-experimental design with comparison group OR pre-post design with sufficient data for effect size computation'),
        ('Data Reporting', 'Sufficient statistical information to compute or convert to Hedges\' g (means and SDs, Cohen\'s d, F, t, r, or pre-post data)'),
        ('Publication Type', 'Published journal articles, conference proceedings, dissertations, or preprints (no publication date restriction)'),
        ('Language', 'English-language full text available')
    ]

    table = doc.add_table(rows=len(criteria)+1, cols=2)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Criterion'
    table.rows[0].cells[1].text = 'Definition'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, (criterion, definition) in enumerate(criteria, 1):
        table.rows[idx].cells[0].text = criterion
        table.rows[idx].cells[1].text = definition

    doc.add_paragraph()
    doc.add_heading('3.2 Exclusion Criteria', level=2)
    doc.add_paragraph('Studies are excluded if they meet ANY of the following:')

    exclusions = [
        'AI system lacks agency characteristics (e.g., static tutoring systems, simple chatbots without adaptivity)',
        'No quantitative learning outcome data reported',
        'No comparison condition (single-group designs without pre-post data)',
        'Insufficient statistical information for effect size computation',
        'Study focuses solely on AI development/technical performance without learning outcomes',
        'Non-educational applications (entertainment, marketing, healthcare without learning component)',
        'Full text not available in English'
    ]

    for exclusion in exclusions:
        doc.add_paragraph(exclusion, style='List Bullet')

    doc.add_page_break()

def add_section_4(doc):
    """Section 4: Effect Size Extraction"""
    doc.add_heading('4. Effect Size Extraction', level=1)

    doc.add_paragraph(
        'All effect sizes are standardized as Hedges\' g with small-sample bias correction. '
        'This section provides computational procedures for extracting effect sizes from various reported statistics.'
    )

    doc.add_heading('4.1 Computing Hedges\' g from Means and SDs', level=2)
    doc.add_paragraph('Preferred method when means and SDs are reported for treatment and control groups:')

    doc.add_paragraph('Step 1: Compute pooled SD:')
    doc.add_paragraph('SD_pooled = sqrt[((n_t - 1) * SD_t² + (n_c - 1) * SD_c²) / (n_t + n_c - 2)]')

    doc.add_paragraph('Step 2: Compute Cohen\'s d:')
    doc.add_paragraph('d = (M_treatment - M_control) / SD_pooled')

    doc.add_paragraph('Step 3: Apply small-sample correction to get Hedges\' g:')
    doc.add_paragraph('J = 1 - (3 / (4 * (n_t + n_c - 2) - 1))')
    doc.add_paragraph('g = d * J')

    doc.add_paragraph('Step 4: Compute standard error:')
    doc.add_paragraph('SE_g = sqrt((n_t + n_c) / (n_t * n_c) + g² / (2 * (n_t + n_c)))')

    doc.add_paragraph('Step 5: Compute 95% confidence interval:')
    doc.add_paragraph('CI_lower = g - 1.96 * SE_g')
    doc.add_paragraph('CI_upper = g + 1.96 * SE_g')

    doc.add_heading('4.2 Converting from Other Statistics', level=2)

    conversions = [
        ('Reported Cohen\'s d', 'd', 'Apply small-sample correction J (see 4.1 Step 3), g = d * J'),
        ('t-statistic', 't', 'd = t * sqrt((n_t + n_c) / (n_t * n_c)), then convert to g'),
        ('F-statistic (df=1)', 'F', 'd = sqrt(F * (n_t + n_c) / (n_t * n_c)), then convert to g'),
        ('Correlation r', 'r', 'd = 2r / sqrt(1 - r²), then convert to g'),
        ('Reported Hedges\' g', 'g', 'Use directly, verify SE computation')
    ]

    table = doc.add_table(rows=len(conversions)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Source Statistic'
    table.rows[0].cells[1].text = 'Symbol'
    table.rows[0].cells[2].text = 'Conversion Formula'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, (source, symbol, formula) in enumerate(conversions, 1):
        table.rows[idx].cells[0].text = source
        table.rows[idx].cells[1].text = symbol
        table.rows[idx].cells[2].text = formula

    doc.add_heading('4.3 Handling Pre-Post Designs', level=2)
    doc.add_paragraph(
        'For pre-post designs without control group, use Morris & DeShon (2002) method for within-subjects designs. '
        'Requires pre-test/post-test means, SDs, and correlation (or assume r = 0.50 if not reported):'
    )

    doc.add_paragraph('d_within = (M_post - M_pre) / SD_pooled_change')
    doc.add_paragraph('SD_pooled_change = sqrt((SD_pre² + SD_post² - 2 * r * SD_pre * SD_post) / 2)')

    doc.add_paragraph('Then apply small-sample correction to obtain g.')

    doc.add_heading('4.4 Multiple Effect Sizes per Study', level=2)
    doc.add_paragraph('When studies report multiple outcomes:')
    items = [
        'Extract ALL effect sizes separately (do not average at extraction)',
        'Assign unique es_id for each effect size (e.g., study001_es01, study001_es02)',
        'Code outcome_type and outcome_measure for each',
        'Statistical dependency will be handled via robust variance estimation (RVE) in analysis',
        'If outcomes are clearly independent domains (e.g., math and reading), code as separate',
        'If outcomes are subscales of same construct, note in discrepancy_notes'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.5 Missing Data Handling', level=2)

    missing_scenarios = [
        ('SD not reported', 'Attempt to compute from SE, CI, t, or F. If impossible, impute from similar studies in same domain (note in es_source).'),
        ('Sample sizes unclear', 'Use reported N, adjust for attrition if reported. Note assumptions in discrepancy_notes.'),
        ('Pre-test correlation unavailable', 'Assume r = 0.50 for pre-post designs (sensitivity analysis will vary this).'),
        ('Control group data missing', 'Exclude effect size if no comparison possible. Note in study-level notes.'),
        ('Incomplete subgroup data', 'Extract only subgroups with complete data. Do not impute missing cells.')
    ]

    table = doc.add_table(rows=len(missing_scenarios)+1, cols=2)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Missing Data Type'
    table.rows[0].cells[1].text = 'Handling Protocol'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, (data_type, protocol) in enumerate(missing_scenarios, 1):
        table.rows[idx].cells[0].text = data_type
        table.rows[idx].cells[1].text = protocol

    doc.add_page_break()

def add_section_5(doc):
    """Section 5: AI Agent Characteristic Coding"""
    doc.add_heading('5. AI Agent Characteristic Coding', level=1)

    doc.add_paragraph(
        'This section provides detailed coding rules for AI agent characteristics. '
        'These codes are critical moderators for RQ2, RQ3, and RQ5.'
    )

    doc.add_heading('5.1 Human Oversight Level', level=2)
    doc.add_paragraph('Codes the degree of human involvement in AI-mediated learning interactions.')

    levels = [
        ('fully_autonomous',
         'AI operates independently with no human intervention during learning sessions. Learner interacts solely with AI. '
         'Teacher/instructor may design initial parameters but does not monitor or intervene.',
         'Example: Autonomous AI tutor providing adaptive practice problems with no teacher access to real-time data.'),

        ('ai_led_checkpoints',
         'AI drives the learning interaction, but includes structured checkpoints where human reviews progress, '
         'provides feedback, or adjusts parameters. Human has oversight role but AI maintains primary instructional control.',
         'Example: AI tutor conducts daily lessons, teacher reviews weekly progress reports and adjusts difficulty settings.'),

        ('human_led_ai_support',
         'Human instructor leads instruction, using AI as supplementary tool. Human makes primary pedagogical decisions; '
         'AI provides data, recommendations, or specific support functions.',
         'Example: Teacher leads classroom instruction, AI provides real-time formative assessment suggestions during lesson.')
    ]

    doc.add_heading('Definitions & Decision Rules:', level=3)
    for code, definition, example in levels:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)
        doc.add_paragraph(example, style='List Bullet')

    doc.add_heading('Coding Decision Tree:', level=3)
    doc.add_paragraph('1. Does a human instructor monitor or intervene during AI interactions? → No: fully_autonomous')
    doc.add_paragraph('2. Who initiates learning activities and makes real-time pedagogical decisions?')
    doc.add_paragraph('   → AI initiates, human reviews periodically: ai_led_checkpoints', style='List Bullet')
    doc.add_paragraph('   → Human initiates, AI supports: human_led_ai_support', style='List Bullet')

    doc.add_heading('5.2 Agent Architecture', level=2)
    doc.add_paragraph('Codes whether the AI system uses single or multiple agents.')

    arch_codes = [
        ('single_agent', 'One AI agent interacts with learner. May have multiple modules/components but operates as unified system.'),
        ('multi_agent', 'Two or more distinct AI agents with different roles/expertise. Agents may collaborate, compete, or specialize.')
    ]

    for code, definition in arch_codes:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_paragraph('For multi_agent systems, code num_agents field with integer count.')

    doc.add_heading('Decision Rules:', level=3)
    items = [
        'Different AI personas with distinct roles → multi_agent',
        'Multiple learners interacting with same AI instance → single_agent',
        'AI with multiple knowledge modules but single interaction interface → single_agent',
        'AI agents simulating peer discussion/debate → multi_agent'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.3 Agency Level (APCP Framework)', level=2)
    doc.add_paragraph(
        'Codes the level of agency based on Yan et al. (2025) APCP framework. '
        'Four levels representing increasing agency capability:'
    )

    apcp_levels = [
        ('adaptive',
         'AI responds to learner input and adapts based on performance or explicit requests. Reactive but personalized.',
         'Characteristics: Performance-based adaptation, content sequencing based on mastery, difficulty adjustment',
         'Example: ITS that increases problem difficulty after consecutive correct answers'),

        ('proactive',
         'AI initiates interactions and anticipates learner needs without explicit prompting. Forward-looking behavior.',
         'Characteristics: Predicts knowledge gaps, suggests next steps, intervenes before errors, sends reminders',
         'Example: AI tutor that detects confusion from hesitation patterns and offers hints before learner requests help'),

        ('co_learner',
         'AI collaborates as learning partner, engaging in joint problem-solving and knowledge construction. Mutual learning stance.',
         'Characteristics: Engages in dialogue, asks clarifying questions, builds on learner ideas, admits uncertainty, co-constructs knowledge',
         'Example: AI partner in collaborative coding project that suggests approaches, debates design decisions, and learns from learner feedback'),

        ('peer',
         'AI operates as equal peer with bidirectional influence, emotional rapport, and social presence. Full social agency.',
         'Characteristics: Emotional recognition/expression, social chat, maintains relationship memory, exhibits personality, mutual influence',
         'Example: AI study buddy that remembers previous conversations, shares study struggles, adjusts interaction style based on learner mood')
    ]

    for code, definition, chars, example in apcp_levels:
        doc.add_heading(code, level=3)
        doc.add_paragraph(definition)
        p = doc.add_paragraph()
        p.add_run('Key Characteristics: ').italic = True
        p.add_run(chars)
        p = doc.add_paragraph()
        p.add_run('Example: ').italic = True
        p.add_run(example)

    doc.add_heading('Coding Decision Process:', level=3)
    doc.add_paragraph('1. Check for social/emotional features → Present: co_learner or peer')
    doc.add_paragraph('2. If social: Does AI maintain relationship memory and exhibit personality? → Yes: peer, No: co_learner')
    doc.add_paragraph('3. If not social: Does AI initiate interactions or predict needs? → Yes: proactive, No: adaptive')

    doc.add_heading('5.4 Agent Role', level=2)

    roles = [
        ('tutor', 'Provides instruction, explanations, scaffolding. Expert-novice dynamic.'),
        ('coach', 'Guides practice, monitors progress, provides encouragement. Performance-oriented.'),
        ('assessor', 'Evaluates knowledge, provides feedback, diagnoses misconceptions.'),
        ('collaborator', 'Works jointly with learner on tasks. Partnership dynamic.'),
        ('facilitator', 'Supports learner-driven exploration, asks guiding questions, manages resources.')
    ]

    for code, definition in roles:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_heading('5.5 Agent Modality', level=2)

    modalities = [
        ('text_only', 'Text-based interaction only (chatbot, messaging interface).'),
        ('voice', 'Speech-based interaction (spoken dialogue, voice assistant).'),
        ('embodied_avatar', 'Visual embodied agent with facial expressions, gestures, or animated character.'),
        ('mixed', 'Combination of modalities (e.g., avatar with text and voice).')
    ]

    for code, definition in modalities:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_heading('5.6 AI Technology', level=2)

    tech = [
        ('rule_based', 'Expert system with predefined rules and decision trees.'),
        ('ML', 'Machine learning models (classification, regression, clustering) without language processing.'),
        ('NLP', 'Natural language processing for understanding/generating text (pre-LLM era techniques).'),
        ('LLM', 'Large language models (GPT, BERT, T5, Claude, etc.).'),
        ('reinforcement_learning', 'RL-based agents learning from interaction feedback.')
    ]

    for code, definition in tech:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_heading('5.7 Adaptivity Level', level=2)

    adaptivity = [
        ('static', 'No adaptation; same content/sequence for all learners.'),
        ('adaptive_performance', 'Adapts based on performance data only (accuracy, response time, mastery).'),
        ('adaptive_behavior_affect', 'Adapts based on behavioral and/or affective data (engagement, emotion, attention, help-seeking).')
    ]

    for code, definition in adaptivity:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_page_break()

def add_section_6(doc):
    """Section 6: Learning Context Coding"""
    doc.add_heading('6. Learning Context Coding', level=1)

    doc.add_paragraph(
        'Learning context variables are key moderators for RQ4. Code the educational setting, '
        'domain, and delivery mode.'
    )

    doc.add_heading('6.1 Learning Context & Education Level', level=2)

    contexts = [
        ('k12', 'Primary/secondary education (ages 5-18)', 'elementary, middle, high_school'),
        ('higher_education', 'Undergraduate or graduate university education', 'undergraduate, graduate'),
        ('workplace_training', 'On-the-job training or corporate learning', 'adult_professional'),
        ('professional_education', 'Certification, licensure, or continuing professional development', 'adult_professional'),
        ('continuing_education', 'Lifelong learning, community education, MOOCs', 'adult_professional')
    ]

    table = doc.add_table(rows=len(contexts)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Learning Context'
    table.rows[0].cells[1].text = 'Definition'
    table.rows[0].cells[2].text = 'Education Level Options'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, (context, definition, levels) in enumerate(contexts, 1):
        table.rows[idx].cells[0].text = context
        table.rows[idx].cells[1].text = definition
        table.rows[idx].cells[2].text = levels

    doc.add_heading('6.2 Domain & Subject', level=2)
    doc.add_paragraph('Broad domain categories:')
    domains = ['STEM (science, technology, engineering, math)', 'language (native or foreign language learning)',
               'medical (healthcare, clinical training)', 'business (management, economics, finance)',
               'ICT (information/communication technology, computer science)', 'arts (visual, performing, design)',
               'other (specify in specific_subject field)']
    for domain in domains:
        doc.add_paragraph(domain, style='List Bullet')

    doc.add_paragraph('Always code specific_subject with detailed topic (e.g., "algebra", "Python programming", "academic writing").')

    doc.add_heading('6.3 Learning Mode & Delivery', level=2)

    modes = [
        ('formal', 'Structured curriculum with credentials (courses, degree programs, certification)'),
        ('informal', 'Self-directed learning without formal credentials (MOOCs, tutorials, apps)'),
        ('blended', 'Mix of formal and informal elements')
    ]

    doc.add_paragraph('Learning Mode:')
    for code, definition in modes:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    delivery = [
        ('face_to_face_with_AI', 'In-person setting with AI tools/agents integrated'),
        ('fully_online', 'Entirely online/remote learning'),
        ('hybrid', 'Mix of in-person and online components')
    ]

    doc.add_paragraph('Delivery Mode:')
    for code, definition in delivery:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_heading('6.4 Comparison Condition', level=2)

    comparisons = [
        ('no_AI', 'Control group receives same instruction/materials without any AI'),
        ('non_agentic_AI', 'Control uses non-agentic AI (e.g., static chatbot, search tool)'),
        ('human_only', 'Human instructor without AI support'),
        ('traditional_instruction', 'Conventional teaching methods (lecture, textbook)')
    ]

    for code, definition in comparisons:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_page_break()

def add_section_7(doc):
    """Section 7: Outcome Classification"""
    doc.add_heading('7. Outcome Classification', level=1)

    doc.add_paragraph(
        'Learning outcomes are classified by type, cognitive level (Bloom\'s taxonomy), '
        'measurement approach, and timing. These codes enable subgroup analysis by outcome type.'
    )

    doc.add_heading('7.1 Outcome Type', level=2)

    outcome_types = [
        ('cognitive', 'Declarative/procedural knowledge (facts, concepts, understanding)', 'Knowledge tests, concept inventories, comprehension assessments'),
        ('skill_based', 'Applied skills and competencies (problem-solving, technical skills)', 'Performance tasks, skill demonstrations, programming assignments'),
        ('affective', 'Attitudes, motivation, self-efficacy, engagement', 'Self-report scales, surveys, behavioral engagement metrics'),
        ('performance', 'Real-world task completion or achievement outcomes', 'Course grades, project quality, certification exam pass rates')
    ]

    table = doc.add_table(rows=len(outcome_types)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Outcome Type'
    table.rows[0].cells[1].text = 'Definition'
    table.rows[0].cells[2].text = 'Example Measures'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, (otype, definition, examples) in enumerate(outcome_types, 1):
        table.rows[idx].cells[0].text = otype
        table.rows[idx].cells[1].text = definition
        table.rows[idx].cells[2].text = examples

    doc.add_heading('7.2 Bloom\'s Taxonomy Level', level=2)
    doc.add_paragraph('Classify cognitive outcomes using revised Bloom\'s taxonomy (3-level aggregation):')

    blooms = [
        ('remember_understand', 'Lower-order: recall, recognition, comprehension', 'Definition tests, multiple-choice recall, summarization'),
        ('apply_analyze', 'Middle-order: application, analysis, problem-solving', 'Problem sets, case analysis, debugging tasks'),
        ('evaluate_create', 'Higher-order: evaluation, synthesis, creation', 'Design projects, critique tasks, original research')
    ]

    for code, definition, examples in blooms:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)
        doc.add_paragraph('Examples: ' + examples, style='List Bullet')

    doc.add_heading('7.3 Measurement Type', level=2)

    measurements = [
        ('standardized_test', 'Published instrument with established psychometrics'),
        ('researcher_developed', 'Custom instrument created for the study'),
        ('performance_assessment', 'Authentic task performance (project, presentation, portfolio)'),
        ('self_report', 'Learner self-assessment or survey')
    ]

    for code, definition in measurements:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_heading('7.4 Timing of Measurement', level=2)

    timings = [
        ('immediate_post', 'Measured immediately after intervention (within 1 week)'),
        ('delayed_post', 'Measured after delay (1+ weeks post-intervention)'),
        ('transfer', 'Measured on transfer tasks (different from training context)')
    ]

    for code, definition in timings:
        p = doc.add_paragraph()
        p.add_run(code).bold = True
        p.add_run(': ' + definition)

    doc.add_page_break()

def add_section_8(doc):
    """Section 8: Quality Assessment"""
    doc.add_heading('8. Quality Assessment', level=1)

    doc.add_paragraph(
        'Risk of bias assessment adapted from Cochrane RoB 2.0 for educational interventions. '
        'Each domain rated as: low_risk, some_concerns, high_risk.'
    )

    domains = [
        ('random_assignment_quality',
         'Randomization procedure and allocation concealment',
         'Low: True randomization with concealed allocation\nSome: Quasi-random (alternating assignment, cohort)\nHigh: Non-random, self-selection'),

        ('sample_representativeness',
         'Sample generalizability and recruitment bias',
         'Low: Representative sample, minimal exclusions\nSome: Convenience sample from typical setting\nHigh: Highly selective sample, volunteer bias'),

        ('measurement_validity',
         'Outcome measure validity and reliability',
         'Low: Standardized, validated instruments\nSome: Researcher-developed with pilot testing\nHigh: Unvalidated measures, single-item indicators'),

        ('attrition_bias',
         'Differential dropout and missing data',
         'Low: <10% attrition, balanced across groups\nSome: 10-20% attrition, no differential pattern\nHigh: >20% attrition or differential dropout'),

        ('reporting_completeness',
         'Completeness of statistical reporting',
         'Low: Complete data for all outcomes, all planned analyses reported\nSome: Minor missing info (e.g., SD imputed)\nHigh: Selective reporting, missing key statistics'),

        ('treatment_fidelity',
         'Implementation adherence and contamination',
         'Low: Monitored implementation, high adherence\nSome: Implementation described but not monitored\nHigh: Poor adherence, contamination between groups')
    ]

    for domain_code, domain_name, criteria in domains:
        doc.add_heading(domain_code, level=2)
        doc.add_paragraph(domain_name)
        p = doc.add_paragraph(criteria)
        p.runs[0].font.size = Pt(10)

    doc.add_heading('8.1 Overall Risk of Bias', level=2)
    doc.add_paragraph('Synthesize domain ratings:')
    items = [
        'Low risk: All domains rated low_risk',
        'Some concerns: At least one domain rated some_concerns, no high_risk domains',
        'High risk: At least one domain rated high_risk'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.2 Sensitivity Analysis Exclusion', level=2)
    doc.add_paragraph(
        'Mark exclude_from_sensitivity = TRUE if overall_risk_of_bias = high_risk. '
        'Sensitivity analyses will compare results with/without high-risk studies.'
    )

    doc.add_page_break()

def add_section_9(doc):
    """Section 9: AI-Assisted Coding Protocol"""
    doc.add_heading('9. AI-Assisted Coding Protocol', level=1)

    doc.add_paragraph(
        'This meta-analysis employs a 7-phase AI-assisted extraction pipeline combining '
        'AI models (GPT-4, Claude Opus) with human oversight for efficiency and accuracy.'
    )

    phases = [
        ('Phase 1: Initial Screening',
         'AI screens titles/abstracts against inclusion criteria. Human reviews AI-flagged borderline cases.',
         'AI: GPT-4 Turbo, Human: 20% double-check'),

        ('Phase 2: Full-Text Eligibility',
         'AI extracts full texts and applies detailed inclusion criteria. Human reviews exclusions.',
         'AI: Claude Opus, Human: 100% exclusion verification'),

        ('Phase 3: Study Metadata Extraction',
         'AI extracts bibliographic data, sample characteristics, study design. Human spot-checks 30%.',
         'AI: GPT-4, Human: 30% random sample'),

        ('Phase 4: Effect Size Extraction',
         'AI locates outcome data, computes effect sizes. Human verifies ALL effect size computations.',
         'AI: Claude Opus (computation), Human: 100% verification'),

        ('Phase 5: AI Agent Coding',
         'AI codes agent characteristics using this manual. Human reviews ALL agency-level codes.',
         'AI: GPT-4 + Claude consensus, Human: 100% review'),

        ('Phase 6: Context & Quality Coding',
         'AI codes learning context and quality assessment. Human reviews 50%.',
         'AI: GPT-4, Human: 50% random sample'),

        ('Phase 7: Discrepancy Resolution',
         'Human expert reconciles all AI-human disagreements. Senior reviewer adjudicates unresolved cases.',
         'Human-led process')
    ]

    table = doc.add_table(rows=len(phases)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Phase'
    table.rows[0].cells[1].text = 'Process'
    table.rows[0].cells[2].text = 'Verification'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, (phase, process, verification) in enumerate(phases, 1):
        table.rows[idx].cells[0].text = phase
        table.rows[idx].cells[1].text = process
        table.rows[idx].cells[2].text = verification

    doc.add_heading('9.1 AI Coding Prompts', level=2)
    doc.add_paragraph(
        'AI models receive structured prompts with this coding manual embedded, study PDF, '
        'and specific extraction task. All AI responses logged with model version, timestamp, '
        'and confidence scores in AI_Extraction_Provenance sheet.'
    )

    doc.add_heading('9.2 Human Verification Protocol', level=2)
    items = [
        'ALL effect size computations verified by human coder',
        'ALL AI agent characteristic codes (especially agency_level_apcp) reviewed by human',
        'RANDOM 30-50% of other codes spot-checked',
        'ALL discrepancies between AI and human flagged for resolution',
        'High-stakes codes (quality assessment, exclusion decisions) receive elevated scrutiny'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

def add_section_10(doc):
    """Section 10: Inter-Coder Reliability"""
    doc.add_heading('10. Inter-Coder Reliability', level=1)

    doc.add_paragraph(
        'Inter-coder reliability assessed on 20% random sample of included studies, '
        'double-coded by independent human coders.'
    )

    doc.add_heading('10.1 Reliability Targets', level=2)

    targets = [
        ('Categorical codes (agent characteristics, context)', 'Cohen\'s κ ≥ 0.80', 'Excellent agreement'),
        ('Effect size extraction', 'ICC(2,1) ≥ 0.90', 'Strong consistency in g values'),
        ('Quality assessment ratings', 'Cohen\'s κ ≥ 0.75', 'Substantial agreement')
    ]

    table = doc.add_table(rows=len(targets)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Variable Type'
    table.rows[0].cells[1].text = 'Statistic'
    table.rows[0].cells[2].text = 'Target Threshold'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for idx, (var_type, stat, threshold) in enumerate(targets, 1):
        table.rows[idx].cells[0].text = var_type
        table.rows[idx].cells[1].text = stat
        table.rows[idx].cells[2].text = threshold

    doc.add_heading('10.2 Coder Training', level=2)
    doc.add_paragraph('All human coders complete:')
    items = [
        'Initial training: Review this manual, practice on 5 pilot studies',
        'Calibration session: Code 3 studies, compare with expert coder, discuss discrepancies',
        'Certification: Achieve κ ≥ 0.80 on 3 consecutive studies before independent coding',
        'Ongoing calibration: Monthly meetings to discuss challenging cases, update manual as needed'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

def add_section_11(doc):
    """Section 11: Discrepancy Resolution"""
    doc.add_heading('11. Discrepancy Resolution', level=1)

    doc.add_paragraph(
        'Systematic process for resolving disagreements between coders or between AI and human coders.'
    )

    doc.add_heading('11.1 Resolution Hierarchy', level=2)
    doc.add_paragraph('Step 1: Independent re-review')
    doc.add_paragraph(
        'Both coders re-examine study independently. If consensus reached, finalize. If not, proceed to Step 2.',
        style='List Bullet'
    )

    doc.add_paragraph('Step 2: Consensus discussion')
    doc.add_paragraph(
        'Coders discuss rationale referencing specific text from study. Attempt to reach consensus using decision rules in this manual.',
        style='List Bullet'
    )

    doc.add_paragraph('Step 3: Senior reviewer adjudication')
    doc.add_paragraph(
        'If no consensus after discussion, senior reviewer (study PI or methodology expert) makes final determination. '
        'Document rationale in discrepancy_notes field.',
        style='List Bullet'
    )

    doc.add_heading('11.2 Priority for Discrepancy Resolution', level=2)
    doc.add_paragraph('High priority (require resolution before analysis):')
    high_priority = [
        'Effect size values (g, SE, CI)',
        'Inclusion/exclusion decisions',
        'Agency level (APCP) classification',
        'Human oversight level',
        'Quality assessment ratings'
    ]
    for item in high_priority:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('Medium priority (resolve during data cleaning):')
    medium_priority = [
        'Learning context codes',
        'Outcome classification',
        'Agent architecture/modality',
        'Sample demographics'
    ]
    for item in medium_priority:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('Low priority (note but may defer):')
    low_priority = [
        'Minor metadata discrepancies (author names, page numbers)',
        'Exact wording of free-text descriptions'
    ]
    for item in low_priority:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('11.3 Documentation', level=2)
    doc.add_paragraph('All discrepancies logged in AI_Extraction_Provenance sheet with:')
    items = [
        'Original AI coding',
        'Human coder disagreement',
        'Resolution outcome',
        'Rationale for final decision',
        'Timestamp of resolution'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

def add_appendix_a(doc):
    """Appendix A: Decision Trees"""
    doc.add_heading('Appendix A: Quick Reference Decision Trees', level=1)

    doc.add_heading('A.1 Human Oversight Level Decision Tree', level=2)
    doc.add_paragraph('START: Does a human monitor or intervene during learning sessions?')
    doc.add_paragraph('  NO → fully_autonomous', style='List Bullet')
    doc.add_paragraph('  YES → Who initiates activities and makes real-time decisions?', style='List Bullet')
    doc.add_paragraph('    AI initiates, human reviews periodically → ai_led_checkpoints', style='List Bullet')
    doc.add_paragraph('    Human initiates, AI supports → human_led_ai_support', style='List Bullet')

    doc.add_heading('A.2 Agency Level (APCP) Decision Tree', level=2)
    doc.add_paragraph('START: Does AI exhibit social/emotional features (relationship memory, personality, empathy)?')
    doc.add_paragraph('  YES → Does it maintain longitudinal relationship and exhibit distinct personality?', style='List Bullet')
    doc.add_paragraph('    YES → peer', style='List Bullet')
    doc.add_paragraph('    NO → co_learner', style='List Bullet')
    doc.add_paragraph('  NO → Does AI initiate interactions or proactively predict needs?', style='List Bullet')
    doc.add_paragraph('    YES → proactive', style='List Bullet')
    doc.add_paragraph('    NO → adaptive', style='List Bullet')

    doc.add_heading('A.3 Effect Size Source Priority', level=2)
    doc.add_paragraph('PREFER (in order):')
    items = [
        '1. Reported Hedges\' g (verify computation)',
        '2. Means and SDs for treatment/control groups',
        '3. Reported Cohen\'s d (apply correction)',
        '4. t-statistic with group sizes',
        '5. F-statistic (df=1) with group sizes',
        '6. Correlation r with sample size',
        '7. Pre-post means/SDs with correlation (or assumed r=0.50)'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('A.4 Multiple Outcomes Handling', level=2)
    doc.add_paragraph('Decision process:')
    items = [
        'Are outcomes measuring same construct? → Extract all, note dependency in discrepancy_notes',
        'Are outcomes independent domains? → Extract separately, code different outcome_type',
        'Are outcomes subscales of single measure? → Extract all, code same outcome_measure with es_id variants',
        'Analysis will use RVE to account for dependency'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

def add_appendix_b(doc):
    """Appendix B: Example Coded Study"""
    doc.add_heading('Appendix B: Example Coded Study', level=1)

    doc.add_paragraph('Hypothetical Example: "AI Peer Tutors in Algebra Learning"')

    doc.add_heading('Study_Metadata Sheet:', level=2)
    metadata_example = {
        'study_id': 'smith2025',
        'authors': 'Smith, J., Lee, K., & Wang, M.',
        'year': '2025',
        'title': 'Effects of Multi-Agent AI Peer Tutoring on Algebra Learning',
        'journal': 'Journal of Educational Technology',
        'doi': '10.1234/jet.2025.001',
        'country': 'USA',
        'n_treatment': '78',
        'n_control': '82',
        'n_total': '160',
        'pct_female': '52',
        'mean_age': '14.5',
        'pub_type': 'journal',
        'study_design': 'RCT',
        'random_assignment': 'yes',
        'duration_weeks': '8',
        'notes': 'Urban middle school, pre-registered trial'
    }

    table = doc.add_table(rows=len(metadata_example), cols=2)
    table.style = 'Light Shading Accent 1'
    for idx, (field, value) in enumerate(metadata_example.items()):
        table.rows[idx].cells[0].text = field
        table.rows[idx].cells[1].text = value
        table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_heading('Effect_Sizes Sheet (excerpt):', level=2)
    es_example = {
        'study_id': 'smith2025',
        'es_id': 'smith2025_es01',
        'outcome_measure': 'Algebra procedural skills test',
        'outcome_type': 'skill_based',
        'outcome_level_blooms': 'apply_analyze',
        'measurement_type': 'standardized_test',
        'timing': 'immediate_post',
        'treatment_m': '78.4',
        'treatment_sd': '12.3',
        'treatment_n': '78',
        'control_m': '71.2',
        'control_sd': '13.1',
        'control_n': '82',
        'hedges_g': '0.56',
        'se_g': '0.16',
        'es_source': 'computed_from_means'
    }

    table = doc.add_table(rows=len(es_example), cols=2)
    table.style = 'Light Shading Accent 1'
    for idx, (field, value) in enumerate(es_example.items()):
        table.rows[idx].cells[0].text = field
        table.rows[idx].cells[1].text = value
        table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_heading('AI_Agent_Characteristics Sheet:', level=2)
    agent_example = {
        'study_id': 'smith2025',
        'agent_name_description': 'Two AI tutors (AlgebraBuddy and MathPal) with distinct personalities',
        'human_oversight_level': 'ai_led_checkpoints',
        'agent_architecture': 'multi_agent',
        'num_agents': '2',
        'agency_level_apcp': 'co_learner',
        'agent_role': 'collaborator',
        'agent_modality': 'text_only',
        'ai_technology': 'LLM',
        'adaptivity': 'adaptive_behavior_affect',
        'feedback_type': 'immediate',
        'personalization': 'both',
        'notes': 'Agents engage in peer discussion, debate solution approaches, teacher reviews weekly dashboards'
    }

    table = doc.add_table(rows=len(agent_example), cols=2)
    table.style = 'Light Shading Accent 1'
    for idx, (field, value) in enumerate(agent_example.items()):
        table.rows[idx].cells[0].text = field
        table.rows[idx].cells[1].text = value
        table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_heading('Coding Rationale:', level=2)
    doc.add_paragraph('Human Oversight: ai_led_checkpoints → AI drives daily sessions, teacher reviews weekly progress.')
    doc.add_paragraph('Agency Level: co_learner → Agents engage in dialogue and debate, but no evidence of emotional rapport or personality persistence.')
    doc.add_paragraph('Architecture: multi_agent → Two distinct AI agents with different personalities collaborating.')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('End of Coding Manual').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    """Generate the Word coding manual"""
    doc = Document()

    # Setup
    setup_styles(doc)

    # Add all sections
    add_title_page(doc)
    add_toc(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)
    add_section_10(doc)
    add_section_11(doc)
    add_appendix_a(doc)
    add_appendix_b(doc)

    # Save
    output_path = "Agentic_AI_Learning_MA_Coding_Manual_v1.docx"
    doc.save(output_path)

    print(f"✓ Created: {output_path}")
    print(f"  - {len(doc.element.body)} elements")
    print(f"  - Sections: Title, TOC, 11 main sections, 2 appendices")

if __name__ == "__main__":
    main()
